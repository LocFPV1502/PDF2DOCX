"""DOCX builder: tạo file Word từ text + OCR results."""

from __future__ import annotations

import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from src.ocr_engine import OcrResult, OcrBlock


def build_docx(
    text_pages: list[Optional[str]],
    ocr_results: dict[int, OcrResult],
    output_path: str,
    source_images: Optional[list[Optional[Image.Image]]] = None,
    include_images: bool = False,
) -> str:
    """Tạo file .docx từ kết quả processing.

    Args:
        text_pages: Text content cho từng trang (None nếu scanned).
        ocr_results: {page_num: OcrResult} cho scanned pages.
        output_path: Đường dẫn file output.
        source_images: Ảnh gốc của scanned pages (để chèn vào DOCX).
        include_images: Có chèn ảnh gốc vào DOCX không.

    Returns:
        Đường dẫn file đã tạo.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    num_pages = len(text_pages)

    for page_num in range(num_pages):
        is_text = text_pages[page_num] is not None
        text = text_pages[page_num]
        ocr = ocr_results.get(page_num)
        img = source_images[page_num] if source_images else None

        # Page break (except first page)
        if page_num > 0:
            doc.add_page_break()

        if is_text:
            # Text layer page → giữ nguyên content
            _add_text_page(doc, text, page_num + 1)
        elif ocr:
            # OCR page → insert text blocks
            _add_ocr_page(doc, ocr, page_num + 1)

            # Optionally chèn ảnh gốc
            if include_images and img:
                _add_image(doc, img)
        elif img:
            # Fallback: có ảnh nhưng không có OCR
            _add_image(doc, img)

    # Save
    doc.save(output_path)
    return output_path


def build_docx_bytes(
    text_pages: list[Optional[str]],
    ocr_results: dict[int, OcrResult],
    source_images: Optional[list[Optional[Image.Image]]] = None,
    include_images: bool = False,
) -> bytes:
    """Tạo file .docx và trả về bytes.

    Dùng khi muốn stream directly đến client.

    Returns:
        DOCX file content as bytes.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        build_docx(
            text_pages=text_pages,
            ocr_results=ocr_results,
            output_path=tmp_path,
            source_images=source_images,
            include_images=include_images,
        )
        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _add_text_page(doc: Document, text: str, page_num: int):
    """Thêm trang text layer vào document."""
    # Add page number header
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"Trang {page_num}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Add text content
    paragraphs = text.split("\n")
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text.strip())


def _add_ocr_page(doc: Document, ocr: OcrResult, page_num: int):
    """Thêm trang OCR content."""
    # Page number header
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"Trang {page_num} (OCR)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Group blocks into paragraphs
    # PaddleOCR trả về blocks đã sort top-to-bottom
    # Group các blocks có Y-coordinate gần nhau thành 1 paragraph
    paragraphs = _group_blocks_to_paragraphs(ocr.blocks)

    for para_lines in paragraphs:
        para_text = " ".join(para_lines)
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

    # Thêm confidence note
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run(
        f"[OCR Confidence: {ocr.average_confidence:.1%}]"
    )
    conf_run.font.size = Pt(8)
    conf_run.font.color.rgb = RGBColor(180, 180, 180)
    conf_run.font.italic = True


def _add_image(doc: Document, img: Image.Image, max_width: float = 5.5):
    """Thêm ảnh vào document."""
    # Convert PIL Image → bytes
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)

    # Calculate size
    width_inches = min(max_width, img.width / 96)  # Assume 96 DPI
    height_inches = width_inches * (img.height / img.width)

    doc.add_picture(buf, width=Inches(width_inches))


def _group_blocks_to_paragraphs(
    blocks: list[OcrBlock],
    y_threshold: float = 20.0,
) -> list[list[str]]:
    """Group OCR blocks thành các dòng paragraph.

    Args:
        blocks: List OcrBlock (đã sort top-to-bottom).
        y_threshold: Khoảng cách Y tối đa để coi là cùng dòng.

    Returns:
        List of paragraphs, mỗi paragraph là list of text lines.
    """
    if not blocks:
        return []

    paragraphs = []
    current_para = [blocks[0].text]
    current_y = blocks[0].bbox[0][1]

    for block in blocks[1:]:
        block_y = block.bbox[0][1]

        # Nếu Y gần nhau → cùng dòng
        if abs(block_y - current_y) <= y_threshold:
            current_para.append(block.text)
        else:
            # Dòng mới
            paragraphs.append(current_para)
            current_para = [block.text]
            current_y = block_y

    # Add last paragraph
    if current_para:
        paragraphs.append(current_para)

    return paragraphs
